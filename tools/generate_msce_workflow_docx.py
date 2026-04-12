# Regenerate DOCX: python3 -m venv .venv && .venv/bin/pip install python-docx && .venv/bin/python tools/generate_msce_workflow_docx.py
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "MSCE_Attendance_App_Workflow_Marathi.docx"


def add_heading_m(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial Unicode MS"
        run.font.size = Pt(14 if level == 1 else 12)
    return p


def add_para(doc: Document, text: str, bullet: bool = False):
    style = "List Bullet" if bullet else None
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.name = "Arial Unicode MS"
        run.font.size = Pt(11)
    return p


def main():
    doc = Document()

    t = doc.add_heading("MSCE Attendance App", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.name = "Arial Unicode MS"
        run.bold = True

    st = doc.add_paragraph("MSCE अटेन्डन्स अॅप — कार्यप्रवाह व वापरकर्ता मार्गदर्शिका")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in st.runs:
        run.font.name = "Arial Unicode MS"
        run.font.size = Pt(12)

    doc.add_paragraph()

    add_heading_m(doc, "१. अॅपची उपलब्धता", 1)
    add_para(
        doc,
        "हे अॅप Android व iOS वापरकर्त्यांसाठी तयार आहे. तुम्ही ते Google Play Store व Apple App Store वरून डाउनलोड करू शकता.",
    )
    add_para(doc, "Google Play Store वरून Android फोनसाठी इंस्टॉल करा.", bullet=True)
    add_para(doc, "Apple App Store वरून iPhone व iPad साठी इंस्टॉल करा.", bullet=True)

    add_heading_m(doc, "२. अॅप एकूणच काय करते?", 1)
    add_para(
        doc,
        "MSCE Attendance App संस्थेतील दैनंदिन हजेरी (उपस्थिती) व्यवस्थित नोंदवण्यास, पडताळण्यास व अहवाल पाहण्यास मदत करते. प्रशासक/शिक्षक सेटअप करतात; विद्यार्थी प्रवेश व निर्गमन फोटो घेऊन हजेरी पूर्ण करतात.",
    )

    add_heading_m(doc, "३. सुरुवातीचा कार्यप्रवाह (संस्था बाजू)", 1)
    add_para(
        doc,
        "बॅच तयार करा: वर्ग किंवा बॅचचे नाव, वर्ष, दिवसाची सुरुवात व शेवटची वेळ, तसेच विषय योग्यरित्या भरा.",
        bullet=True,
    )
    add_para(
        doc,
        "विद्यार्थी जोडा: प्रत्येक विद्यार्थ्यास योग्य बॅच द्या. बॅच निवडल्यावर त्या वेळापत्रकानुसार हजेरीचे नियम लागू होतात. नोंदणीसाठी चेहऱ्याचा स्पष्ट फोटो घ्या.",
        bullet=True,
    )
    add_para(
        doc,
        "GPS / स्थान सेटिंग: संस्थेचे ठिकाण अॅपमध्ये सेट करा. हजेरीसाठी विद्यार्थ्यांना संस्थेच्या ठिकाणाजवळ (अॅपने ठरवलेल्या अंतरापर्यंत) असणे गरजेचे असते.",
        bullet=True,
    )

    add_heading_m(doc, "४. विद्यार्थी हजेरी कशी नोंदवतात?", 1)
    add_para(
        doc,
        "प्रवेश फोटो: व्याख्यान सुरू झाल्यापासून ठराविक मर्यादेत (उदा. सुरुवातीच्या २० मिनिटांत) प्रवेश फोटो घेता येतो.",
        bullet=True,
    )
    add_para(
        doc,
        "प्रवेश फोटो घेतल्यावर त्या दिवसाच्या वेळापत्रकातील व्याख्यानांसाठी उपस्थिती दर्शविली जाते.",
        bullet=True,
    )
    add_para(
        doc,
        "निर्गमन फोटो: शेवटच्या व्याख्यानाच्या आधी व नंतरच्या वेळेत निर्गमन फोटो घेता येतो. निर्गमनाने प्रवेश ते निर्गमन दरम्यानची उपस्थिती पुष्ट होते.",
        bullet=True,
    )
    add_para(
        doc,
        "जर प्रवेश घेतला पण निर्गमन विंडोमध्ये निर्गमन फोटो घेतला नाही, तर अॅपच्या नियमांनुसार संबंधित व्याख्यानांसाठी अनुपस्थित मानले जाऊ शकते.",
        bullet=True,
    )
    add_para(
        doc,
        "चेहरा ओळख व लाइव्ह फोटो: फक्त थेट घेतलेला फोटो स्वीकारला जातो; फोटोचा फोटो, स्क्रीनशॉट किंवा छापील चित्र नाकारले जाऊ शकते.",
        bullet=True,
    )

    add_heading_m(doc, "५. हजेरी मोजणी व आकडे — स्क्रीनवर काय दिसते?", 1)
    add_para(
        doc,
        "प्रत्येक व्याख्यानासाठी स्थिती: उपस्थित किंवा अनुपस्थित — अहवालात तसेच दिसते.",
        bullet=True,
    )
    add_para(
        doc,
        "प्रवेश वेळ, निर्गमन वेळ व प्रत्येक व्याख्यानाची तपशीलवार स्थिती अहवालात पाहता येते.",
        bullet=True,
    )
    add_para(
        doc,
        "दैनंदिन आकडेवारी: आज किती विद्यार्थी उपस्थित आहेत, एकूण विद्यार्थ्यांपैकी टक्केवारी इत्यादी सारांश दृश्यांमध्ये दिसू शकते.",
        bullet=True,
    )
    add_para(
        doc,
        "कॅलेंडर व ट्रेंड: दिवस, आठवडा किंवा महिन्यानुसार नमुने व कमी हजेरीचे विश्लेषण पाहता येते.",
        bullet=True,
    )
    add_para(
        doc,
        "PDF: एका किंवा सर्व विद्यार्थ्यांचे अहवाल PDF म्हणून चढवून रेकॉर्ड ठेवता येतात.",
        bullet=True,
    )

    add_heading_m(doc, "६. सुरक्षा व विश्वासार्हता (वापरकर्ता दृष्टीने)", 1)
    add_para(
        doc,
        "चेहरा ओळख: योग्य विद्यार्थ्याची ओळख पडताळली जाते.",
        bullet=True,
    )
    add_para(
        doc,
        "स्थान तपासणी: हजेरी संस्थेच्या जवळूनच घेता येईल अशा प्रकारे डिझाइन केलेले आहे.",
        bullet=True,
    )
    add_para(
        doc,
        "वेळ विंडो: वर्गाच्या बाहेरच्या वेळेत हजेरी टाळण्यास वेळेचे नियम आहेत.",
        bullet=True,
    )
    add_para(
        doc,
        "लाइव्हनेस: थेट व्यक्ती असल्याची खात्री करण्यास मदत — फसवणूक कमी करणे.",
        bullet=True,
    )

    add_heading_m(doc, "७. चांगल्या निकासासाठी टिपा", 1)
    add_para(
        doc,
        "चांगला प्रकाश, स्थिर कॅमेरा व चेहरा स्पष्ट दिसावा; डोळे उघडे ठेवा; मास्क/गॉगलमुळे ओळख अयशस्वी होऊ शकते.",
        bullet=True,
    )
    add_para(
        doc,
        "स्थान परवानग्या व GPS सुरू ठेवा जेणेकरून स्थान तपासणी पूर्ण व्हावी.",
        bullet=True,
    )
    add_para(
        doc,
        "इंटरनेट कमी असताना काही नोंदी नंतर सिंक होऊ शकतात; अॅपमधील सिंक स्थिती लक्षात ठेवा.",
        bullet=True,
    )

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "दस्तऐवज: MSCE Attendance App — फक्त अॅप व वापरकर्ता कार्यप्रवाह; तांत्रिक बॅकएंड/डेटाबेस तपशील यात समाविष्ट नाहीत."
    )
    for run in foot.runs:
        run.font.name = "Arial Unicode MS"
        run.font.size = Pt(9)
        run.italic = True

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
